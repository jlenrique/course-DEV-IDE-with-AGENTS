"""Unit tests for source_wrangler_operations."""

from __future__ import annotations

import importlib.util
import io
import json
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as _DocxDocument
from docx.shared import Inches
from PIL import Image
from pypdf import PdfWriter

_MODULE_DIR = Path(__file__).resolve().parents[1]
_REPO_ROOT = _MODULE_DIR.parents[2]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

_SPEC = importlib.util.spec_from_file_location(
    "source_wrangler_operations",
    _MODULE_DIR / "source_wrangler_operations.py",
)
assert _SPEC and _SPEC.loader
sw = importlib.util.module_from_spec(_SPEC)
sys.modules["source_wrangler_operations"] = sw
_SPEC.loader.exec_module(sw)


def test_html_to_text_strips_tags() -> None:
    html = "<html><body><p>Hello <b>world</b></p><script>x</script></body></html>"
    assert "Hello" in sw.html_to_text(html)
    assert "world" in sw.html_to_text(html)
    assert "x" not in sw.html_to_text(html)


def test_build_extracted_markdown() -> None:
    md = sw.build_extracted_markdown(
        "Trial",
        [("Section A", "body a"), ("Section B", "body b")],
    )
    assert "# Source bundle: Trial" in md
    assert "## Section A" in md
    assert "body a" in md


def test_write_source_bundle(tmp_path: Path) -> None:
    prov = [
        sw.SourceRecord(kind="url", ref="https://example.com", note="test"),
    ]
    summary = sw.write_source_bundle(
        tmp_path / "b1",
        "My bundle",
        "# extracted\n",
        prov,
        raw_files={"sample.html": "<p>x</p>"},
    )
    assert Path(summary["extracted_md"]).exists()
    meta = json.loads((tmp_path / "b1" / "metadata.json").read_text())
    assert meta["title"] == "My bundle"
    assert (tmp_path / "b1" / "raw" / "sample.html").exists()


def test_wrangle_playwright_saved_html(tmp_path: Path) -> None:
    p = tmp_path / "capture.html"
    p.write_text("<html><body><h1>Title</h1><p>Para</p></body></html>", encoding="utf-8")
    title, text, rec = sw.wrangle_playwright_saved_html(p, source_url="https://gamma.app/x")
    assert rec.kind == "playwright_html"
    assert rec.ref == "https://gamma.app/x"
    assert "Para" in text


def test_is_gamma_app_docs_url() -> None:
    assert sw.is_gamma_app_docs_url("https://gamma.app/docs/abc?mode=doc")
    assert sw.is_gamma_app_docs_url("https://www.gamma.app/docs/x")
    assert not sw.is_gamma_app_docs_url("https://public-api.gamma.app/v1.0/themes")
    assert not sw.is_gamma_app_docs_url("https://example.com/")


def test_fetch_url_rejects_gamma_docs() -> None:
    with pytest.raises(sw.GammaDocsURLNotSupportedError):
        sw.fetch_url("https://gamma.app/docs/test-id")


@patch.object(sw.requests, "get")
def test_summarize_url_for_envelope(mock_get: MagicMock) -> None:
    mock_resp = MagicMock()
    mock_resp.content = b"<html><body><p>OK</p></body></html>"
    mock_resp.headers = {"Content-Type": "text/html; charset=utf-8"}
    mock_resp.encoding = "utf-8"
    mock_resp.raise_for_status = MagicMock()
    mock_get.return_value = mock_resp
    _title, text, rec = sw.summarize_url_for_envelope("https://example.org/doc")
    assert rec.kind == "url"
    assert "OK" in text


def test_list_box_files(tmp_path: Path) -> None:
    (tmp_path / "a.md").write_text("x", encoding="utf-8")
    (tmp_path / "sub").mkdir()
    (tmp_path / "sub" / "b.txt").write_text("y", encoding="utf-8")
    found = sw.list_box_files(tmp_path, max_files=20)
    assert len(found) == 2


def test_verify_and_require_local_sources(tmp_path: Path) -> None:
    existing = tmp_path / "a.txt"
    existing.write_text("ok", encoding="utf-8")
    missing = tmp_path / "gone.pdf"
    assert sw.verify_local_source_paths([existing]) == []
    assert sw.verify_local_source_paths([missing]) == [missing]
    sw.require_local_source_files([existing])
    with pytest.raises(FileNotFoundError, match="gone"):
        sw.require_local_source_files([missing])


def test_wrangle_local_pdf_blank_page(tmp_path: Path) -> None:
    pdf_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with pdf_path.open("wb") as f:
        writer.write(f)
    title, text, rec = sw.wrangle_local_pdf(pdf_path)
    assert title == "blank"
    assert rec.kind == "local_pdf"
    assert "pypdf" in rec.note
    assert isinstance(text, str)


def test_extract_pdf_text_sample_deck() -> None:
    sample = _REPO_ROOT / "course-content" / "staging" / "Diagnosis-Innovation.pdf"
    if not sample.is_file():
        pytest.skip("staging sample PDF not present")
    body, meta = sw.extract_pdf_text(sample, max_pages=2)
    assert meta["engine"] == "pypdf"
    assert meta["pages_total"] >= 1
    assert "Innovation" in body or "innovation" in body.lower()


# ---------------------------------------------------------------------------
# Story 27-1: DOCX provider wiring — unit tests
#
# Fixtures generated in-test via python-docx Document() builder; no binary
# .docx is committed to the repo. Per AC-T.4: pillow PNG uses io.BytesIO
# (NOT a tmp_path file) to avoid parallel-pytest disk race.
# ---------------------------------------------------------------------------


def _build_happy_path_docx(path: Path) -> None:
    """Synthetic DOCX with 3 heading levels, 2 tables, a numbered list, and
    ~300 words. Tables are interleaved between paragraphs to catch
    iteration-order drift (AC-T.1 Murat-guard)."""
    doc = _DocxDocument()
    doc.add_heading("Chapter One: Introduction", level=1)
    doc.add_paragraph(
        "This sample document exercises the wrangle_local_docx extraction path. "
        "It contains multi-sentence paragraphs, nested headings, tables that "
        "interleave with body text, and a numbered list. " * 2
    )
    # First table, between heading and next paragraph block
    t1 = doc.add_table(rows=2, cols=3)
    t1.rows[0].cells[0].text = "Col A"
    t1.rows[0].cells[1].text = "Col B"
    t1.rows[0].cells[2].text = "Col C"
    t1.rows[1].cells[0].text = "r1c1"
    t1.rows[1].cells[1].text = "r1c2"
    t1.rows[1].cells[2].text = "r1c3"

    doc.add_heading("Section 1.1: Middle Content", level=2)
    doc.add_paragraph(
        "Additional narrative text follows the first table. The extractor must "
        "render this paragraph after the table, preserving document order. " * 2
    )
    # Numbered list (python-docx renders as ListNumber style)
    for i in range(1, 5):
        doc.add_paragraph(f"Numbered item {i}: short entry text.", style="List Number")

    doc.add_heading("Subsection 1.1.1: Deeper", level=3)
    doc.add_paragraph(
        "Trailing narrative after the second heading level. The fixture ensures "
        "all three heading levels are exercised by extraction. " * 2
    )
    # Second table, at the tail
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "K"
    t2.rows[0].cells[1].text = "V"

    doc.save(str(path))


def _build_empty_docx(path: Path) -> None:
    """DOCX with zero author-added paragraphs. python-docx's Document() still
    creates the root body with one empty paragraph by default — the extractor
    must treat this as empty content gracefully (AC-T.3)."""
    doc = _DocxDocument()
    doc.save(str(path))


def _build_images_only_docx(path: Path) -> None:
    """DOCX containing only an inline PNG image, no text runs (AC-T.4).
    PNG built in-memory via io.BytesIO; NOT written to tmp_path to avoid
    parallel-pytest disk-file race."""
    png_buffer = io.BytesIO()
    img = Image.new("RGB", (32, 32), color=(200, 100, 50))
    img.save(png_buffer, format="PNG")
    png_buffer.seek(0)

    doc = _DocxDocument()
    doc.add_picture(png_buffer, width=Inches(1.0))
    doc.save(str(path))


def test_wrangle_local_docx_happy_path(tmp_path: Path) -> None:
    """AC-T.1: Happy-path extraction preserves structure.

    Synthetic DOCX with 3 heading levels, 2 interleaved tables, numbered list,
    ~300 words. Asserts (title, body, rec) shape + key content presence.
    """
    docx_path = tmp_path / "chapter_one.docx"
    _build_happy_path_docx(docx_path)

    title, body, rec = sw.wrangle_local_docx(docx_path)

    # Title mirrors wrangle_local_pdf convention: stem with underscores → spaces
    assert title == "chapter one"
    assert rec.kind == "local_docx"
    # Note string reports counts (shape: "python-docx paragraphs=N headings=M tables=K")
    assert "python-docx" in rec.note
    assert "paragraphs=" in rec.note
    assert "headings=" in rec.note
    assert "tables=" in rec.note

    # Body content must contain:
    # - Heading 1 rendered as `# Chapter One: Introduction`
    assert "# Chapter One: Introduction" in body
    # - Heading 2 rendered as `## Section 1.1: Middle Content`
    assert "## Section 1.1: Middle Content" in body
    # - Heading 3 rendered as `### Subsection 1.1.1: Deeper`
    assert "### Subsection 1.1.1: Deeper" in body
    # - Paragraph text (sample phrase from first paragraph)
    assert "multi-sentence paragraphs" in body
    # - Table 1 pipe-row rendering
    assert "| Col A | Col B | Col C |" in body
    assert "| r1c1 | r1c2 | r1c3 |" in body
    # - Table 2 pipe-row rendering
    assert "| K | V |" in body
    # - Numbered list items (text preserved; style marker not required)
    assert "Numbered item 1" in body
    assert "Numbered item 4" in body

    # Body-order discipline: first heading appears BEFORE first table in body
    assert body.index("# Chapter One: Introduction") < body.index("| Col A | Col B | Col C |")
    # Second heading appears AFTER first table
    assert body.index("| r1c1 | r1c2 | r1c3 |") < body.index("## Section 1.1: Middle Content")


def test_wrangle_local_docx_corrupt_file_raises(tmp_path: Path) -> None:
    """AC-T.2a: Library-layer corrupt DOCX raises a caller-catchable exception.

    Scope: wrangle_local_docx raises when python-docx cannot open the file.
    Fixture: raw bytes with .docx suffix that fail the ZIP-magic check
    (malformed-ZIP branch specifically, per Winston's green-light note —
    NOT a valid-ZIP-wrong-content fixture which is a different failure mode).
    """
    bad_path = tmp_path / "broken.docx"
    bad_path.write_bytes(b"not a docx file at all")

    with pytest.raises(Exception) as exc_info:
        sw.wrangle_local_docx(bad_path)
    # python-docx surfaces PackageNotFoundError for malformed ZIP
    exc_name = type(exc_info.value).__name__
    exc_msg = str(exc_info.value)
    assert "PackageNotFoundError" in exc_name or "Package" in exc_msg


def test_wrangle_local_docx_empty_document(tmp_path: Path) -> None:
    """AC-T.3: Empty DOCX (zero author-added paragraphs) extracts without crash.

    python-docx creates a root body with one empty paragraph by default;
    extractor must treat this gracefully: zero headings, zero tables, near-empty body.
    """
    docx_path = tmp_path / "empty.docx"
    _build_empty_docx(docx_path)

    title, body, rec = sw.wrangle_local_docx(docx_path)

    assert title == "empty"
    assert rec.kind == "local_docx"
    assert "headings=0" in rec.note
    assert "tables=0" in rec.note
    # Body is empty or contains only whitespace / a single empty-paragraph marker
    assert len(body.strip()) <= 1


def test_wrangle_local_docx_images_only(tmp_path: Path) -> None:
    """AC-T.4: Images-only DOCX (inline PNG, no text) extracts without crash.

    Asserts no text content extracted; note string reports zero paragraphs/
    headings; no traceback from missing text runs. Pillow PNG is built via
    io.BytesIO (NOT a tmp_path file) to avoid parallel-pytest disk race.
    """
    docx_path = tmp_path / "images_only.docx"
    _build_images_only_docx(docx_path)

    title, body, rec = sw.wrangle_local_docx(docx_path)

    assert title == "images only"
    assert rec.kind == "local_docx"
    assert "headings=0" in rec.note
    assert "tables=0" in rec.note
    # Text body should be empty (or near-empty — docx may carry an image-anchor paragraph
    # whose text is empty string)
    assert body.strip() == ""


def test_wrangle_source_maps_docx_corruption_to_failed_outcome(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """AC-T.2b: Adapter-layer translates DOCX corruption to FAILED SourceOutcome.

    Scope: _wrangle_source() via run_wrangler catches the DOCX extractor's
    exception and maps to FAILED outcome with error_kind=docx_extraction_failed
    and known_losses=[docx_open_failed]. Atomicity rule (Murat): library-raise
    and adapter-mapping fail for different reasons; red-light message must
    point at the right layer.

    Also verifies AC-B.3's "no traceback leaks to stdout" requirement via
    capsys (code-review Acceptance Auditor, 2026-04-17).
    """
    # Locate run_wrangler via the same load_module_from_path pattern run_wrangler itself uses
    from scripts.utilities.skill_module_loader import load_module_from_path

    run_wrangler = load_module_from_path(
        "texas_run_wrangler_under_test",
        _MODULE_DIR / "run_wrangler.py",
    )

    bad_path = tmp_path / "broken.docx"
    bad_path.write_bytes(b"definitely not a docx")

    src = {
        "ref_id": "test_corrupt_docx",
        "provider": "local_file",
        "locator": str(bad_path),
        "role": "primary",
        "description": "corrupt DOCX under test",
    }
    now = "2026-04-17T22:00:00+00:00"

    outcome = run_wrangler._wrangle_source(src, now)

    assert outcome.error_kind == "docx_extraction_failed"
    # error_detail is "{ExceptionClass}: {message}" per existing _wrangle_source pattern
    assert outcome.error_detail is not None
    detail = outcome.error_detail
    assert "PackageNotFoundError" in detail or "Package not found" in detail
    # known_losses MUST be the DOCX-specific sentinel (AC-B.3), not the generic template
    assert outcome.report.known_losses == ["docx_open_failed"]
    # Tier is FAILED (no content extracted)
    assert outcome.report.tier.name == "FAILED"
    assert outcome.content_text == ""
    # AC-B.3: "No traceback leaks to stdout" — adapter must swallow the
    # exception cleanly. Capsys captures stdout + stderr; neither should
    # contain a Python traceback header.
    captured = capsys.readouterr()
    assert "Traceback (most recent call last)" not in captured.out
    assert "Traceback (most recent call last)" not in captured.err
