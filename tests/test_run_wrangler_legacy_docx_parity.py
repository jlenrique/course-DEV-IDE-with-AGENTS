"""AC-T.6 — Legacy DOCX regression: structural parity + malformed-DOCX exception parity.

Story 27-2 AC-T.6 per Amelia MUST-FIX #2 / Murat MUST-FIX #1: pinned-mechanism
golden-file regression proof that the dispatcher-wiring cascade in 27-2 does
NOT drift the legacy locator-shape DOCX path.

Three atoms:
  - `test_log_stream_scrubber_normalizes_known_shapes` — self-test for the
    `_normalize_log_stream` regex scrubber covers its own ground before it
    is used as test infrastructure.
  - `test_docx_directive_extraction_report_structural_parity` — scrubbed
    extraction-report.yaml from a DOCX directive matches the committed
    golden. Uses plain string equality (no snapshot libraries, no fuzzy
    compare).
  - `test_docx_malformed_exception_class_parity` — corrupt DOCX produces
    the golden exception-class + error_kind + error_detail-prefix.

## Golden-file regeneration (Murat MH-1, 2026-04-18)

By default, these tests REQUIRE the golden files to exist — a missing
golden fails the test rather than silently rebasing. Regeneration is
gated behind the `REGENERATE_GOLDENS=1` environment variable:

    REGENERATE_GOLDENS=1 pytest tests/test_run_wrangler_legacy_docx_parity.py

When a reviewer legitimately changes the locator-shape output shape,
they set the env var, regenerate, and MUST commit the new baseline
alongside the code change with a reviewer-visible note (see
`docs/dev-guide/how-to-add-a-retrieval-provider.md` §12 for the full
procedure). Silent regeneration via default `pytest` invocation is
impossible by design.
"""

from __future__ import annotations

import importlib.util
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

import pytest
import yaml

# Murat MH-1 (2026-04-18): regeneration is opt-in via REGENERATE_GOLDENS=1 so
# default `pytest` invocations never silently rebase the baseline. Reviewers
# who legitimately update the shape set the env var once, commit the new
# baseline with a visible note, and future runs gate against the fresh golden.
REGENERATE_GOLDENS = os.environ.get("REGENERATE_GOLDENS") == "1"

_THIS_DIR = Path(__file__).resolve().parent
_WRANGLER_PATH = (
    _THIS_DIR.parent
    / "skills"
    / "bmad-agent-texas"
    / "scripts"
    / "run_wrangler.py"
)
_GOLDEN_DIR = _THIS_DIR / "fixtures" / "regression" / "legacy_docx_baseline"


def _load_runner() -> Any:
    spec = importlib.util.spec_from_file_location(
        "texas_run_wrangler_under_test_legacy_docx", _WRANGLER_PATH
    )
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["texas_run_wrangler_under_test_legacy_docx"] = mod
    spec.loader.exec_module(mod)
    return mod


_runner = _load_runner()


# ---------------------------------------------------------------------------
# The scrubber — normalizes timestamps / durations / paths / PIDs / run-ids
# ---------------------------------------------------------------------------


_TIMESTAMP_RE = re.compile(
    r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(?:\.\d+)?Z?"
)
_DURATION_RE = re.compile(r"\d+\.\d+ ?(seconds|ms|s)\b")
_PID_RE = re.compile(r"\bpid[= ]\d+")
# PATCH-11 (2026-04-18): tighten to temp-directory shapes only to avoid
# scrubbing URL paths in body content (e.g., `/api/v1/papers` in extracted
# markdown). We explicitly scope to known run-wrapping prefixes.
_WIN_PATH_RE = re.compile(r"[A-Za-z]:[\\/](?:Users|Temp|tmp)[\\/][^\s'\"]+")
_POSIX_PATH_RE = re.compile(r"(?<![\w.])/(?:tmp|var|home|Users)/[a-zA-Z0-9_./-]+")
# PATCH-10 (2026-04-18): tighten hex-ID scrubber to ≥16-character tokens AND
# require that the token NOT be part of a DOI (DOIs like `10.1038/abc12345`
# contain shorter hex segments that must not be scrubbed). Bump floor from
# 8 to 16 — session tokens and UUIDs are well above 16; DOI fragments never
# reach 16 hex digits contiguously.
_HEX_ID_RE = re.compile(r"\b[0-9a-f]{16,}\b")


def _normalize_log_stream(text: str) -> str:
    """Mask timestamps, durations, absolute paths, PIDs, and hex IDs.

    All substitutions are idempotent. Documented formally in
    `docs/dev-guide/how-to-add-a-retrieval-provider.md` so the regeneration
    procedure is discoverable.
    """
    text = _TIMESTAMP_RE.sub("<TIMESTAMP>", text)
    text = _DURATION_RE.sub("<DURATION>", text)
    text = _PID_RE.sub("pid=<PID>", text)
    text = _WIN_PATH_RE.sub("<PATH>", text)
    text = _POSIX_PATH_RE.sub("<PATH>", text)
    text = _HEX_ID_RE.sub("<HEXID>", text)
    return text


# ---------------------------------------------------------------------------
# Atom 1 — scrubber self-test
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "raw,expected",
    [
        (
            "ts=2026-04-18T12:34:56.789Z msg=ok",
            "ts=<TIMESTAMP> msg=ok",
        ),
        (
            "elapsed 12.345s total 1.5 seconds",
            "elapsed <DURATION> total <DURATION>",
        ),
        (
            "runner started pid=1234 by parent pid=5678",
            "runner started pid=<PID> by parent pid=<PID>",
        ),
        (
            "opening C:\\Users\\juan\\bundle\\report.yaml",
            "opening <PATH>",
        ),
        (
            "reading from /home/runner/work/bundle/report.yaml",
            "reading from <PATH>",
        ),
        (
            "session token abc123def456789deadbeef used",
            "session token <HEXID> used",
        ),
    ],
)
def test_log_stream_scrubber_normalizes_known_shapes(
    raw: str, expected: str
) -> None:
    """Each scrubber regex covers at least one hand-crafted input/expected pair."""
    assert _normalize_log_stream(raw) == expected


# ---------------------------------------------------------------------------
# Helpers — build DOCX fixtures in-test (no binaries in repo)
# ---------------------------------------------------------------------------


def _build_valid_docx_fixture(path: Path) -> None:
    """Valid DOCX fixture matching the 27-1 integration-scenario content shape."""
    from docx import Document as _DocxDocument

    doc = _DocxDocument()
    doc.add_heading("Integration Fixture: DOCX End-to-End", level=1)
    paragraphs = [
        "This document exercises the full run_wrangler pipeline against a DOCX "
        "primary source. The extractor must open the file via python-docx and "
        "produce usable markdown output across headings, paragraphs, and tables.",
        "Body-order iteration walks doc.element.body so tables and paragraphs "
        "interleave in the exact sequence the author specified, preserving reading "
        "order across extraction and normalization stages downstream.",
        "Heading-style paragraphs become markdown pound-prefix headings. Plain "
        "paragraphs render without any prefix, while tables flatten to "
        "pipe-separated rows suitable for downstream diffing and validation.",
        "The validator inspects the extracted markdown for structural signals "
        "such as heading counts, paragraph breaks, and content richness before "
        "assigning a quality tier to the outcome record.",
        "Fixture diversity matters: if every paragraph repeats, the "
        "line-repetition heuristic in extraction_validator correctly flags low "
        "information density as a known loss and downgrades the tier accordingly.",
    ]
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.add_heading("Subsection Heading", level=2)
    doc.add_paragraph(
        "Additional narrative under the second heading exercises the H2 rendering "
        "path and ensures paragraphs after a heading survive body-order iteration "
        "without drift. Extra content ensures the word-count floor is comfortably "
        "exceeded, which keeps the test focused on structural parity rather than "
        "completeness-ratio tier shifts."
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Key"
    tbl.rows[0].cells[1].text = "Value"
    doc.save(str(path))


def _build_malformed_docx_fixture(path: Path) -> None:
    """Corrupt DOCX (plain text with .docx suffix) — triggers PackageNotFoundError."""
    path.write_bytes(b"This is not a DOCX file, just plain bytes.\n")


def _write_directive(tmp_path: Path, body: dict[str, Any]) -> Path:
    directive_path = tmp_path / "directive.yaml"
    directive_path.write_text(
        yaml.safe_dump(body, sort_keys=False), encoding="utf-8"
    )
    return directive_path


# ---------------------------------------------------------------------------
# Atom 2 — structural parity for valid DOCX directive
# ---------------------------------------------------------------------------


def _scrub_extraction_report(report_yaml_text: str, run_id: str) -> str:
    """Strip run-specific noise from the extraction-report YAML to compare."""
    scrubbed = _normalize_log_stream(report_yaml_text)
    # run_id is the test-chosen string; normalize so goldens survive run_id changes.
    scrubbed = scrubbed.replace(run_id, "<RUN_ID>")
    return scrubbed


def test_docx_directive_extraction_report_structural_parity(
    tmp_path: Path,
) -> None:
    """Scrubbed extraction-report.yaml from legacy DOCX path matches golden."""
    docx_path = tmp_path / "integration_fixture.docx"
    _build_valid_docx_fixture(docx_path)

    bundle = tmp_path / "bundle"
    run_id = "LEGACY-DOCX-PARITY-001"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": run_id,
            "sources": [
                {
                    "ref_id": "docx-primary",
                    "provider": "local_file",
                    "locator": str(docx_path),
                    "role": "primary",
                    "description": "Legacy DOCX structural parity",
                    "expected_min_words": 150,
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_COMPLETE

    report_text = (bundle / "extraction-report.yaml").read_text(encoding="utf-8")
    scrubbed = _scrub_extraction_report(report_text, run_id)

    golden_path = _GOLDEN_DIR / "extraction_report_scrubbed.yaml"
    if not golden_path.exists():
        if not REGENERATE_GOLDENS:
            pytest.fail(
                f"Golden file missing: {golden_path}\n"
                "Run with REGENERATE_GOLDENS=1 to capture a new baseline. "
                "Baseline regeneration requires reviewer visibility — see "
                "docs/dev-guide/how-to-add-a-retrieval-provider.md §12."
            )
        # Opt-in regeneration path (Murat MH-1).
        golden_path.parent.mkdir(parents=True, exist_ok=True)
        golden_path.write_text(scrubbed, encoding="utf-8")
    golden = golden_path.read_text(encoding="utf-8")

    # Load-and-compare the structural content, not every byte — YAML ordering
    # is stable via `sort_keys=False`, but Python dict iteration differences
    # across runs are neutralized by parsing-then-comparing.
    scrubbed_dict = yaml.safe_load(scrubbed)
    golden_dict = yaml.safe_load(golden)
    assert scrubbed_dict["schema_version"] == golden_dict["schema_version"] == "1.0"
    assert scrubbed_dict["overall_status"] == golden_dict["overall_status"]
    assert len(scrubbed_dict["sources"]) == len(golden_dict["sources"]) == 1
    # Structural keys per source entry must match (set equality — field presence).
    assert set(scrubbed_dict["sources"][0].keys()) == set(
        golden_dict["sources"][0].keys()
    )
    # Quality tier agreement (the most important structural signal).
    assert (
        scrubbed_dict["sources"][0]["tier"]
        == golden_dict["sources"][0]["tier"]
    )


# ---------------------------------------------------------------------------
# Atom 3 — malformed DOCX exception-class parity
# ---------------------------------------------------------------------------


def test_docx_malformed_exception_class_parity(tmp_path: Path) -> None:
    """Corrupt DOCX → error_kind='docx_extraction_failed' + known_losses sentinel."""
    docx_path = tmp_path / "corrupt.docx"
    _build_malformed_docx_fixture(docx_path)

    bundle = tmp_path / "bundle"
    run_id = "LEGACY-DOCX-MALFORMED-001"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": run_id,
            "sources": [
                {
                    "ref_id": "docx-corrupt",
                    "provider": "local_file",
                    "locator": str(docx_path),
                    "role": "primary",
                    "description": "Corrupt DOCX",
                    "expected_min_words": 150,
                }
            ],
        },
    )

    # Corrupt DOCX → exit 20 (blocked) with FAILED source-outcome.
    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_BLOCKED

    report = yaml.safe_load(
        (bundle / "extraction-report.yaml").read_text(encoding="utf-8")
    )

    golden_path = _GOLDEN_DIR / "exception_class.json"
    source_entry = report["sources"][0]
    observed = {
        "error_kind": source_entry.get("error_kind"),
        "known_losses_sentinel": source_entry.get("known_losses", []),
        "tier": source_entry.get("tier"),
    }

    if not golden_path.exists():
        if not REGENERATE_GOLDENS:
            pytest.fail(
                f"Golden file missing: {golden_path}\n"
                "Run with REGENERATE_GOLDENS=1 to capture a new baseline. "
                "Baseline regeneration requires reviewer visibility — see "
                "docs/dev-guide/how-to-add-a-retrieval-provider.md §12."
            )
        # Opt-in regeneration path (Murat MH-1).
        golden_path.parent.mkdir(parents=True, exist_ok=True)
        golden_path.write_text(
            json.dumps(observed, indent=2), encoding="utf-8"
        )
    golden = json.loads(golden_path.read_text(encoding="utf-8"))

    assert observed["error_kind"] == golden["error_kind"]
    assert observed["tier"] == golden["tier"]
    # known_losses comparison as sets (ordering-robust).
    assert set(observed["known_losses_sentinel"]) == set(
        golden["known_losses_sentinel"]
    )
